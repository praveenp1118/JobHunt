"""
CV file parser — extracts raw text from PDF, DOCX, or plain text.
Output is always a clean string ready for Claude to convert to markdown.
"""
import io
from typing import Tuple


async def parse_file_to_text(file_bytes: bytes, content_type: str, filename: str) -> Tuple[str, str]:
    """
    Parse uploaded file to raw text.
    Returns (raw_text, detected_format).
    """
    filename_lower = filename.lower()

    if content_type == "application/pdf" or filename_lower.endswith(".pdf"):
        return await _parse_pdf(file_bytes), "pdf"

    if (
        content_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        )
        or filename_lower.endswith((".docx", ".doc"))
    ):
        return await _parse_docx(file_bytes), "docx"

    if content_type in ("text/plain", "text/markdown") or filename_lower.endswith((".txt", ".md")):
        return file_bytes.decode("utf-8", errors="replace"), "text"

    # Fallback: try PDF then DOCX
    try:
        return await _parse_pdf(file_bytes), "pdf"
    except Exception:
        pass
    try:
        return await _parse_docx(file_bytes), "docx"
    except Exception:
        pass

    raise ValueError(f"Unsupported file format: {filename}")


async def _parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pypdf."""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())
    full_text = "\n\n".join(pages)
    if not full_text.strip():
        raise ValueError("PDF appears to be scanned/image-based — no text extracted")
    return full_text


async def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)

    return "\n".join(lines)


def count_words(text: str) -> int:
    """Count words in text."""
    return len(text.split())
